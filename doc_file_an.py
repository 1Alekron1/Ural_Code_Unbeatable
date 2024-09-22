import re

from docx import Document
from docx.shared import RGBColor
from test import analyse_text, check_text


def check_structural_elements(doc):
    required_sections = [
        "Титульный лист", "Список исполнителей", "Реферат",
        "Содержание", "Термины и определения",
        "Перечень сокращений и обозначений", "Введение",
        "Основная часть отчета о НИР", "Заключение"
    ]

    found_sections = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().lower()
        for section in required_sections:
            if section.lower() in text:
                found_sections.append(section)

    missing_sections = set(required_sections) - set(found_sections)
    return missing_sections, found_sections


def extract_references_list(doc):
    # Ищем последнее вхождение раздела "Литература" или "Список литературы"
    references = []
    last_reference_section_index = -1

    # Ищем последний индекс вхождения раздела "Литература"
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().lower()
        if "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ".lower() in text or "список литературы" in text:
            last_reference_section_index = i

    # Если раздел найден, собираем пункты списка литературы
    if last_reference_section_index != -1:
        for paragraph in doc.paragraphs[last_reference_section_index + 1:]:
            # Проверяем, является ли абзац пунктом списка (есть ли нумерация)
            if paragraph._element.xpath(".//w:numPr"):
                references.append(paragraph)
            elif paragraph.text.strip() == "":
                # Если встретилась пустая строка, считаем конец списка
                break

    return references


def check_references_in_text(doc, references):
    # Поиск всех сносок в квадратных скобках
    references_in_text = []
    reference_pattern = re.compile(r'\[(\d+)\]')

    for paragraph in doc.paragraphs:
        matches = reference_pattern.findall(paragraph.text)
        if matches:
            references_in_text.extend([int(m) for m in matches])

    # Проверяем индексацию
    is_indexed_correctly = references_in_text == sorted(references_in_text)

    # Проверка на использование всех ссылок из списка литературы
    used_references = set(references_in_text)
    all_references = set(range(1, len(references) + 1))
    unused_references = all_references - used_references

    return references_in_text, is_indexed_correctly, unused_references


def highlight_unused_references(doc, references, unused_references):
    for i, paragraph in enumerate(references):
        ref_num = i + 1
        if ref_num in unused_references:
            # Выделяем красным неиспользованные ссылки
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 0, 0)  # Устанавливаем красный цвет текста


def highlight_incorrectly_indexed_references(doc, references_in_text):
    # Поиск всех сносок и их индексации
    sorted_references = sorted(references_in_text)

    for paragraph in doc.paragraphs:
        text = paragraph.text
        matches = re.finditer(r'\[(\d+)\]', text)
        for match in matches:
            ref_number = int(match.group(1))
            # Если текущая ссылка не на месте (не соответствует правильному порядку)
            if ref_number != sorted_references.pop(0):
                # Подсвечиваем эту ссылку
                start_idx = match.start(1)
                end_idx = match.end(1)
                for run in paragraph.runs:
                    if run.text in match.group(0):  # Подсвечиваем весь текст ссылки
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет для неправильной ссылки


def extract_conclusion_text(doc):
    # Ищем раздел "Заключение"
    conclusion_start_index = -1
    conclusion_text = []

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().lower()
        if "заключение" == text:
            conclusion_start_index = i
            break

    # Если раздел "Заключение" найден, извлекаем текст до следующего раздела
    if conclusion_start_index != -1:
        for paragraph in doc.paragraphs[conclusion_start_index + 1:]:
            text = paragraph.text.strip()
            # Проверяем, не начался ли следующий раздел (например, другой заголовок)
            if text and text.isupper():
                break  # Следующий раздел найден
            conclusion_text.append(paragraph.text)

    return "\n".join(conclusion_text)


def extract_section_text(doc, section_name):
    section_start_index = -1
    section_text = []

    # Поиск начала указанного раздела
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().lower()
        if section_name.lower() == text:
            section_start_index = i
            break

    # Если раздел найден, извлекаем текст до следующего заголовка
    if section_start_index != -1:
        for paragraph in doc.paragraphs[section_start_index + 1:]:
            text = paragraph.text.strip()
            # Проверяем, не начался ли следующий раздел (например, заглавными буквами)
            if text and text.isupper():
                break  # Следующий раздел найден
            section_text.append(paragraph.text)

    return "\n".join(section_text)  # Собираем текст раздела


def check(doc_path):
    entry_criteria = """1) оценка современного состояния решаемой научно-технической проблемы
        2)  основание и исходные данные для разработки темы 
        3) обоснование необходимости проведения НИР
        4) сведения о планируемом научно-техническом уровне разработки, о патентных исследованиях и выводы из них
        5)  актуальность и новизна темы
        6) связь данной работы с другими научно-исследовательскими работами.
        7) цель и задачи исследования"""
    conclusion_criteria = """1) краткие выводы по результатам выполненной НИР или отдельных ее этапов;
    2) оценку полноты решений поставленных задач;
3) разработку рекомендаций и исходных данных по конкретному использованию результатов НИР;
4) результаты оценки технико-экономической эффективности внедрения;
5) результаты оценки научно-технического уровня выполненной НИР в сравнении с лучшими достижениями в этой области."""

    result = {"Введение": "", "Заключение": "", "Пересечение Введения и Заключения по смыслу": "", "Найдены разделы": [], "Отсутствуют разделы": [],
              "Найденные ссылки на литературу": [], "Найденные ссылки в тексте": [], "Ошибка": "",
              "Неиспользованные ссылки из списка литературы": []}
    doc = Document(doc_path)
    conclusion = extract_section_text(doc, "Заключение")
    entry = extract_section_text(doc, "Введение")
    result["Введение"] = analyse_text(entry, entry_criteria)
    result["Заключение"] = analyse_text(conclusion, conclusion_criteria)
    result["Пересечение Введения и Заключения по смыслу"] = check_text(entry, conclusion)

    # Проверка структурных элементов
    missing_sections, found_sections = check_structural_elements(doc)
    result["Найдены разделы"] = found_sections
    result["Отсутствуют разделы"] = missing_sections

    # Извлечение списка литературы
    reference_paragraphs = extract_references_list(doc)
    references = [p.text for p in reference_paragraphs]
    result["Найденные ссылки на литературу"] = references

    # Проверка ссылок в тексте
    references_in_text, is_indexed_correctly, unused_references = check_references_in_text(doc, references)
    result["Найденные ссылки в тексте"] = references_in_text
    if is_indexed_correctly:
        result["Ошибка"] = "Индексация ссылок правильная."
    else:
        result["Ошибка"] = "Ошибка в индексации ссылок."
        # Подсвечиваем ссылки, которые идут в неправильном порядке

    if unused_references:
        result["Неиспользованные ссылки из списка литературы"] = unused_references
        highlight_unused_references(doc, reference_paragraphs, unused_references)

    # Сохранение документа с примечаниями
    doc.save("Исходник с правками.docx")
    for i in result:
        print(f"{i}:\n{result[i]}")
    return result


# Использование:
doc_path = "Бордунов Александр Максимович Отчет по работе.docx"